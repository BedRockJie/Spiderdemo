'''
文档的   doucument 文档
        Pargraph  段落
        Run   文字块  文字的标识不同
'''
from docx import Document
doc = Document('一些操作.docx')

# for pargraph in doc.paragraphs:
#     print(pargraph.text)
paragraph  =  doc.paragraphs[1]
runs = paragraph.runs
for run in runs:
    print(run.text)

