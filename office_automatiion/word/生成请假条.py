from docx import Document
import time
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document() #生成一个文件对象
doc.add_heading('请假条',0)

name = input('请输入姓名：')
why = input('请输入请假原因：')
days = input('请输入请假天数：')
day = int(days)

p = doc.add_paragraph('今有')
p.add_run(name).bold = True
p.add_run('因')
p.add_run(why).bold= True
p.add_run('请假')
p.add_run(days)
p.add_run('天，望批准。')


time1 = time.strftime("%Y-%m-%d")
year = time.strftime("%Y")
mounrh = time.strftime("%m")
day1 = str(int(time.strftime("%d"))+day)

doc.add_paragraph('请假日期：' + time1 + '    '+'请假截至日期' + "{}-{}-{}".format(year,mounrh,day1))

paragraph = doc.add_paragraph('请假人：'+name)
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

doc.save("请假条生成.docx")


# document.add_page_break()  #添加分页
##########段落缩进###########
# paragraph_format = paragraph.paragraph_format  #创建段落格式对象
# paragraph_format.left_indent = Inches(0.5)  #段落左缩进0.5英寸
# #需要  from docx.shared import Inches
# paragraph_format.right_indent = Pt(20)   #右缩进20点
# #from docx.shared import Pt
# paragraph_format.first_line_indent = Inches(0.5)  #第一行缩进

# 可以通过一个Run对象的font属性来获取和设置该Run的字符格式
# 下画线格式font.underline（True表示单下画线，False表示没有下画线