from docx import Document
from docx.shared import Cm
# doc = Document('一些操作.docx')
# # doc.add_heading('一级标题',level=1)
# paragraph1 = doc.add_paragraph('这是一个段落')
# paragraph2 = doc.add_paragraph('这又是一个段落')
#
# paragraph1.add_run('加粗').bold = True


document = Document()
# paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')


document.add_heading('这是标题',0)
p = document.add_paragraph('段落添加到这里')
p.add_run('粗体').bold = True # 加粗
p.add_run('普通')
p.add_run('斜体').italic = True # 斜体


assert isinstance(document.add_page_break, object)
document.add_page_break() #添加分页
document.add_picture(r'D:\正点原子linux\高清壁纸\高清壁纸\1.jpg',width = Cm(19)) #添加图片  按照比例压缩
document.add_table(rows=5,cols=5)
records = [
    [1,2,3,4,5], [1,2,3,4,5], [1,2,3,4,5], [1,2,3,4,5], [1,2,3,4,5],
]
table = document.add_table(rows=5,cols=5)
for row in range(5):
    cells = table.rows[row].cells
    for col in range(5):
        cells[col].text = str(records[row][col])

# document.add_heading('Heading, level 1', level=1)
document.save('demp.docx') # 最后必须点保存